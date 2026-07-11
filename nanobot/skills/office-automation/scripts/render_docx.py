"""Render a report DSL into a docx file."""

from __future__ import annotations

import argparse
import os
from pathlib import Path
from typing import Any

from _common import load_facts, read_json, render_text_value, replace_fact_placeholders
from docx import Document
from officecli_backend import render_with_officecli


def _add_metrics_table(document: Document, items: list[Any], facts: dict[str, dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Metric"
    header[1].text = "Value"
    for item in items:
        if not isinstance(item, dict):
            continue
        row = table.add_row().cells
        row[0].text = str(item.get("label", item.get("fact_ref", "")))
        row[1].text = render_text_value({"fact_ref": item.get("fact_ref")}, facts)


def _add_table(document: Document, block: dict[str, Any], facts: dict[str, dict[str, Any]]) -> None:
    headers = block.get("headers", [])
    rows = block.get("rows", [])
    if not isinstance(headers, list) or not isinstance(rows, list) or not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for source_row in rows:
        if not isinstance(source_row, list):
            continue
        row = table.add_row().cells
        for index, cell in enumerate(source_row[: len(headers)]):
            row[index].text = render_text_value(cell, facts)


def render_docx_legacy(
    *,
    dsl_path: Path,
    facts_path: Path,
    output_path: Path,
    template_path: Path | None = None,
) -> None:
    dsl = read_json(dsl_path)
    facts = load_facts(facts_path)
    document = Document(str(template_path)) if template_path else Document()

    title = replace_fact_placeholders(str(dsl.get("title", "Report")), facts)
    document.add_heading(title, level=0)
    subtitle = dsl.get("subtitle")
    if isinstance(subtitle, str) and subtitle.strip():
        document.add_paragraph(replace_fact_placeholders(subtitle, facts))

    for section in dsl.get("sections", []):
        if not isinstance(section, dict):
            continue
        document.add_heading(replace_fact_placeholders(str(section.get("title", "")), facts), level=1)
        for block in section.get("blocks", []):
            if not isinstance(block, dict):
                continue
            block_type = block.get("type")
            if block_type == "paragraph":
                document.add_paragraph(replace_fact_placeholders(str(block.get("text", "")), facts))
            elif block_type == "bullets":
                for item in block.get("items", []):
                    text = item.get("text") if isinstance(item, dict) else item
                    document.add_paragraph(
                        replace_fact_placeholders(str(text), facts),
                        style="List Bullet",
                    )
            elif block_type == "metrics":
                items = block.get("items", [])
                if isinstance(items, list):
                    _add_metrics_table(document, items, facts)
            elif block_type == "table":
                _add_table(document, block, facts)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def render_docx(
    *,
    dsl_path: Path,
    facts_path: Path,
    output_path: Path,
    template_path: Path | None = None,
    backend: str = "officecli",
    officecli_bin: str | None = None,
    allow_unverified_officecli: bool = False,
    preview_dir: Path | None = None,
) -> None:
    if backend == "python":
        render_docx_legacy(
            dsl_path=dsl_path,
            facts_path=facts_path,
            output_path=output_path,
            template_path=template_path,
        )
        return

    if template_path is not None:
        raise ValueError(
            "OfficeCLI template rendering is intentionally deferred to the dump/merge workflow; "
            "use --backend python for the legacy template path."
        )
    render_with_officecli(
        "docx",
        dsl_path=dsl_path,
        facts_path=facts_path,
        output_path=output_path,
        binary=officecli_bin,
        allow_unverified_version=allow_unverified_officecli,
        preview_dir=preview_dir,
    )


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--dsl", dest="dsl_path", required=True, type=Path)
    parser.add_argument("--facts", dest="facts_path", required=True, type=Path)
    parser.add_argument("--template", dest="template_path", type=Path)
    parser.add_argument("--out", dest="output_path", required=True, type=Path)
    parser.add_argument(
        "--backend",
        choices=("officecli", "python"),
        default=os.environ.get("MYBOT_OFFICE_BACKEND", "officecli"),
    )
    parser.add_argument("--officecli-bin")
    parser.add_argument("--allow-unverified-officecli", action="store_true")
    parser.add_argument("--preview-dir", type=Path)
    args = parser.parse_args()

    render_docx(
        dsl_path=args.dsl_path,
        facts_path=args.facts_path,
        output_path=args.output_path,
        template_path=args.template_path,
        backend=args.backend,
        officecli_bin=args.officecli_bin,
        allow_unverified_officecli=args.allow_unverified_officecli,
        preview_dir=args.preview_dir,
    )


if __name__ == "__main__":
    main()
